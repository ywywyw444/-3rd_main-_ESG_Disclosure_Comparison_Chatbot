# kor_chatbot.py – KSSB 및 IFRS 국문 전용 ESG 챗봇

import os
import re
import streamlit as st
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.messages import SystemMessage, HumanMessage
from langchain.retrievers.multi_query import MultiQueryRetriever
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMListwiseRerank

# ✅ 환경 변수 로드
load_dotenv()

# ✅ 모델 정의
llm = ChatOpenAI(model="gpt-4o", temperature=0)
embedding_model = OpenAIEmbeddings()

# ✅ 클린 텍스트 함수
def clean_text(text: str) -> str:
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^\d{1,2}$", stripped):
            continue
        cleaned.append(stripped)
    return ' '.join(cleaned)

# ✅ 벡터DB 로딩
kor_vector_paths = {
    "IFRS_S1": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\KOKRvectorstores\KOKRvectorstores\KOR_IFRS_S1_2nd",
    "IFRS_S2": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\KOKRvectorstores\KOKRvectorstores\KOR_IFRS_S2_2nd",
    "KSSB1": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\KOKRvectorstores\KOKRvectorstores\KSSB_01_2nd",
    "KSSB2": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\KOKRvectorstores\KOKRvectorstores\KSSB_02_2nd",
    "KSSB101": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\KOKRvectorstores\KOKRvectorstores\KSSB_101_2nd",
}

kor_dbs = {
    name: FAISS.load_local(path, embedding_model, allow_dangerous_deserialization=True)
    for name, path in kor_vector_paths.items()
}

# ✅ 기준 자동 분류
def classify_kor_standards(query: str):
    prompt = f"""
다음 질문을 보고 관련 있는 기준서들을 골라줘. 중복 선택 가능해.
사용자 질문: "{query}"
선택 가능한 기준: IFRS_S1, IFRS_S2, KSSB1, KSSB2, KSSB101
리스트 형식으로 정확히 반환해줘. 예시: ["IFRS_S1", "KSSB2"]
"""
    response = llm([
        SystemMessage(content="너는 ESG 기준서 분류 전문가야."),
        HumanMessage(content=prompt)
    ])
    try:
        return eval(response.content)
    except:
        return list(kor_dbs.keys())

# ✅ 기준별 문서 검색 및 응답 생성
def process_kor_standard(standard, db, query_ko):
    try:
        multi_retriever = MultiQueryRetriever.from_llm(
            retriever=db.as_retriever(search_kwargs={"k": 4}),
            llm=llm
        )
        reranker = LLMListwiseRerank.from_llm(llm=llm, top_n=2)
        compressor = ContextualCompressionRetriever(
            base_retriever=multi_retriever,
            base_compressor=reranker
        )
        docs = compressor.invoke(query_ko)
        if not docs:
            return (standard, None, [])

        context_blocks = []
        references = []
        for doc in docs:
            meta = doc.metadata
            source = meta.get("source", "")
            citation = f"[출처: {standard} - {source or '해당 문단'}]"
            context_blocks.append(f"{citation}\n{doc.page_content}")
            references.append({
                "standard": standard,
                "source": source or "알 수 없음",
                "content": doc.page_content
            })

        context = "\n\n".join(context_blocks)
        system_msg = SystemMessage(
            content=f"{standard} 기준서 전문가로서, 아래 문서를 바탕으로 한국어 질문에 답해주세요:\n\n[문서]\n{context}\n\n[질문]\n{query_ko}"
        )
        user_msg = HumanMessage(content=query_ko)
        response = llm([system_msg, user_msg])
        return (standard, response.content, references)

    except Exception as e:
        print(f"❌ 오류 발생 - {standard}: {e}")
        return (standard, None, [])

# ✅ 기준별 응답 통합
def generate_kor_responses(query_ko: str, selected_standards: list):
    responses = {}
    references = {}

    with ThreadPoolExecutor() as executor:
        futures = [
            executor.submit(process_kor_standard, std, kor_dbs[std], query_ko)
            for std in selected_standards if std in kor_dbs
        ]
        for future in futures:
            std, resp, refs = future.result()
            if resp:
                responses[std] = resp
                references[std] = refs
            else:
                print(f"⛔ 관련 문서 없음: {std}, 응답 제외됨")

    return responses, references

# ✅ Streamlit 실행 함수
def run_kor_chatbot_app():
    import os
    import re
    import streamlit as st
    from io import BytesIO
    import pythoncom
    from docx import Document
    import tempfile
    from docx2pdf import convert

    def remove_control_characters(text: str) -> str:
        return re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

    # st.set_page_config(page_title="ESG 공시 비교 챗봇", layout="wide")
    st.title("🌺K-Mate")

    user_input = st.text_input("질문을 입력하세요 (예: 온실가스 배출을 어떻게 공시해야 하나요?)")

    if st.button("입력") and user_input:
        with st.spinner("질문 분석 중: 관련 기준서를 식별하는 중입니다..."):
            selected = classify_kor_standards(user_input)
            st.markdown(f"✅ 선택된 기준서: {', '.join(selected) if selected else '없음'}")

        with st.spinner("기준별 문서 검색 및 응답 생성 중..."):
            responses, references = generate_kor_responses(user_input, selected)
        
        st.session_state.responses = responses
        st.session_state.references = references
        st.session_state.user_input = user_input
         
    if "user_input" in st.session_state:
        responses = st.session_state.responses
        references = st.session_state.references
        user_input = st.session_state.user_input

        # st.header(f"📄 기준별 상세 응답 보기")
        # for std, resp in responses.items():
        #     st.markdown(f"### {std} 응답")
        #     st.write(resp)

        if responses:
            st.header("📄 기준별 상세 응답 보기")
            for std, resp in responses.items():
                st.markdown(f"### {std} 응답")
                st.write(resp)
        else:
            st.header("📄 기준별 상세 응답 없음")

        with st.expander("📚 인용된 문서 출처 보기"):
            for std, refs in references.items():
                st.markdown(f"## {std} 문서 인용")
                for ref in refs:
                    cleaned_text = clean_text(ref['content'])
                    content_html = f"""
                    <div style='margin-bottom:1.5em; padding:1em; border:1px solid #ccc; border-radius:10px; background-color:#f9f9f9'>
                        <p><b>📌 소스:</b> {ref.get("source", "없음")}</p>
                        <pre style='white-space: pre-wrap; font-size: 14px; color: auto; line-height: 1.5;'>{cleaned_text}</pre>
                    </div>
                    """
                    st.markdown(content_html, unsafe_allow_html=True)

        # Word 문서 생성
        doc = Document()
        doc.add_heading('ESG 공시 비교 챗봇 결과', 0)
        doc.add_heading('질문', level=1)
        doc.add_paragraph(user_input)
        doc.add_heading('📄 기준별 응답', level=1)
        for std, resp in responses.items():
            doc.add_heading(std, level=2)
            doc.add_paragraph(resp)
        doc.add_heading('📚 인용 문서', level=1)
        for std, refs in references.items():
            doc.add_heading(std, level=2)
            for ref in refs:
                doc.add_paragraph(f"출처: {ref['source']}")
                safe_text = remove_control_characters(clean_text(ref['content']))
                doc.add_paragraph(safe_text)

        # Word → BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        st.download_button(
            label="⬇️ Word 파일 다운로드",
            data=doc_io,
            file_name="esg_답변기록.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # PDF 다운로드
        def convert_docx_to_pdf(doc: Document) -> BytesIO:
            pythoncom.CoInitialize()
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "temp.docx")
                pdf_path = os.path.join(tmpdir, "temp.pdf")
                doc.save(docx_path)
                convert(docx_path, pdf_path)
                with open(pdf_path, "rb") as f:
                    return BytesIO(f.read())

        # ✅ PDF 다운로드 버튼 처리
        try:
            pdf_bytes = convert_docx_to_pdf(doc)
            st.download_button(
                label="⬇️ PDF 파일 다운로드",
                data=pdf_bytes,
                file_name="esg_답변기록.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.warning("⚠️ PDF 저장 중 오류가 발생했습니다. Word 파일은 정상 저장됩니다.")
            st.text(str(e))
