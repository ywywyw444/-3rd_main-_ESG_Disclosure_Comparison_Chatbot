# ✅ 기준 분류 기반으로 선택된 기준서만 처리하도록 리팩토링한 코드

import os
import re
import streamlit as st
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.messages import SystemMessage, HumanMessage
from langchain.retrievers.multi_query import MultiQueryRetriever
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMListwiseRerank

# ✅ 환경변수 로딩
load_dotenv()

# ✅ 모델 정의
llm = ChatOpenAI(model="gpt-4o", temperature=0)
embedding_model = OpenAIEmbeddings()

# ✅ 클린업 함수
def clean_text(text: str) -> str:
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^\d{1,2}$", stripped):
            continue
        cleaned.append(stripped)
    return ' '.join(cleaned)

# ✅ 기준서별 벡터 DB 로드
vectorstore_paths = {
    "GRI": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\vectorstores\GRI_01_2nd",
    "IFRS_S1": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\vectorstores\IFRS_01_2nd",
    "IFRS_S2": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\vectorstores\IFRS_02_2nd",
    "TCFD": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\vectorstores\TFCD_01_2nd",
    "KSSB1": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\vectorstores\KSSB_01_2nd",
    "KSSB2": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\vectorstores\KSSB_02_2nd",
    "KSSB101": r"D:\bit_esg\python\esg_chatbot_project\src\vector_finish\vectorstores\KSSB_101_2nd"
}

dbs = {
    name: FAISS.load_local(path, embedding_model, allow_dangerous_deserialization=True)
    for name, path in vectorstore_paths.items()
}

# ✅ GPT 번역 함수
def translate_to_english(korean_query: str) -> str:
    messages = [
        SystemMessage(content="You are a professional English translator. " \
        "Translate the following Korean question into natural and accurate English."),
        HumanMessage(content=korean_query)
    ]
    return llm(messages).content.strip()

# ✅ 기준 키워드 분류 함수
# def extract_relevant_sentences(text, keywords, max_sentences=3):
#     """
#     주어진 응답에서 keywords와 연관된 핵심 문장만 추출
#     """
#     sentences = re.split(r'(?<=[.!?])\s+', text)
#     scored = []

#     for sent in sentences:
#         score = sum(kw.lower() in sent.lower() for kw in keywords)
#         if score > 0:
#             scored.append((score, sent))

#     # 점수 순 정렬 후 상위 max_sentences만 반환
#     top_sentences = [s for _, s in sorted(scored, key=lambda x: -x[0])[:max_sentences]]
#     return " ".join(top_sentences)


# ✅ 기준 분류 함수

def classify_relevant_standards(query: str):
    prompt = f"""
사용자 질문: "{query}"

이 질문과 관련 있는 ESG 기준서를 아래 중에서 선택해줘. 복수 선택 가능해.
가능한 선택: GRI, IFRS_S1, IFRS_S2, TCFD, KSSB1, KSSB2, KSSB101
형식: [\"GRI\", \"IFRS_S2\"] 처럼 Python 리스트로만 답해줘.
"""
    response = llm([
        SystemMessage(content="ESG 기준서 라우팅 전문가"),
        HumanMessage(content=prompt)
    ])
    try:
        return eval(response.content)
    except:
        return list(dbs.keys())


def translate_to_korean(english_text: str) -> str:
    messages = [
        SystemMessage(content="You are a professional Korean translator. Translate the following English ESG-related paragraph into accurate Korean."),
        HumanMessage(content=english_text)
    ]
    return llm(messages).content.strip()

def is_comparative_question_via_llm(query: str) -> bool:
    messages = [
        SystemMessage(content="너는 ESG 공시 질문을 분석하는 전문가야. 사용자 질문이 '기준서 간 비교'를 요구하는지 판단해줘."),
        HumanMessage(content=f"""
사용자 질문: "{query}"

이 질문이 기준서(GRI, IFRS, KSSB 등) 간 비교를 요청하는 질문이면 'True'라고만 대답해줘.  
비교가 아닌 단일 기준서 설명이면 'False'라고만 대답해줘.
""")
    ]
    try:
        response = llm(messages).content.strip()
        return response.lower() == "true"
    except:
        return False




# ✅ 기준별 응답 처리 함수
def process_standard(standard, db, query_ko, query_en):
    try:
        multi_query_retriever = MultiQueryRetriever.from_llm(
            retriever=db.as_retriever(search_kwargs={"k": 4}),
            llm=llm
        )
        reranker = LLMListwiseRerank.from_llm(llm=llm, top_n=2)
        compression_retriever = ContextualCompressionRetriever(
            base_retriever=multi_query_retriever,
            base_compressor=reranker
        )
        docs = compression_retriever.invoke(query_en)

        if not docs:
            return (standard, None, [])

        context_blocks = []
        references = []
        for doc in docs:
            meta = doc.metadata
            source = meta.get("source", "")
            citation = f"[출처: {standard} - {source or '해당 문단'}]"
            context_blocks.append(f"{citation}\n{doc.page_content}")
            references.append({
                "standard": standard,
                "source": source or "알 수 없음",
                "content": doc.page_content
            })

        context = "\n\n".join(context_blocks)
        system_msg = SystemMessage(
            # content=f"You are a professional ESG expert specialized in {standard} reporting standards. Based on the documents below, answer the Korean user query.\n\n[Documents]\n{context}\n\n[Original Korean Query]\n{query_ko}"
        content=f"""You are a professional ESG expert specialized in {standard} reporting standards.
                    You must ONLY use the following documents to answer the Korean query.
                    Do NOT add any external knowledge or assumptions beyond what is explicitly stated in the documents.



                    [Documents]
                    {context}

                    [Original Korean Query]
                    {query_ko}
                    """)
        user_msg = HumanMessage(content=query_ko)
        response = llm([system_msg, user_msg])
        return (standard, response.content, references)

    except Exception as e:
        print(f"❌ 오류 발생 - {standard}: {e}")
        return (standard, None, [])

# ✅ 기준별 응답 생성 (선택 기준만)
# ✅ 기준별 응답 생성 (선택 기준만)
def generate_response_for_each_standard(query_ko: str, selected_standards: list):
    query_en = translate_to_english(query_ko)
    responses = {}
    documents = {}

    with ThreadPoolExecutor() as executor:
        futures = [
            executor.submit(process_standard, std, dbs[std], query_ko, query_en)
            for std in selected_standards if std in dbs
        ]
        for future in futures:
            std, resp, refs = future.result()
            if resp:
                responses[std] = resp
                documents[std] = refs
            else:
                print(f"⛔ 관련 문서 없음: {std}, 응답 제외됨")

    # ✅ 기준별 응답 비교 요약 생성 (단일 문장 형태)
    try:
        summary_prompt = f"""
다음은 사용자 질문에 대해 각 ESG 기준서(GRI, IFRS 등)로부터 수집한 응답입니다.
이 응답들을 비교하여 공통점 또는 차이점을 간결하게 한 문장으로 요약해줘.

[질문]
{query_ko}

[기준별 응답들]
"""
        for std, resp in responses.items():
            summary_prompt += f"\n### {std} 응답:\n{resp}\n"

        messages = [
            SystemMessage(content="너는 ESG 기준서 비교 전문가야. 아래 기준별 응답을 비교한 요약을 단 한 문장으로 말해줘."),
            HumanMessage(content=summary_prompt)
        ]
        summary_sentence = llm(messages).content.strip()
        responses["summary"] = summary_sentence  # ✅ 여기에 최종 문장 삽입
    except Exception as e:
        print(f"⚠️ 요약 생성 실패: {e}")
        responses["summary"] = "⚠️ 기준서 간 요약을 생성하는 데 실패했습니다."

    return responses, documents

import difflib

def is_similar_sentence(a: str, b: str, threshold: float = 0.8) -> bool:
    return difflib.SequenceMatcher(None, a, b).ratio() > threshold


# # ✅ 기준별 요약 응답 생성
# def generate_comparative_summary(responses_dict: dict, original_query_ko: str):
#     filtered_responses = {
#         std: resp for std, resp in responses_dict.items()
#         if "관련 내용 없음" not in resp and resp.strip()
#     }

#     # ✅ 질의에서 키워드 추출 (단순하게 띄어쓰기 기준 분리 또는 tokenizer 적용)
#     keywords = original_query_ko.strip().split()  # 또는 Okt/nltk 등 사용 가능

#     prompt = f"""
# [사용자 질문]
# {original_query_ko}

# [응답 지시사항]
# - 아래 각 기준({', '.join(filtered_responses.keys())})에 대한 응답 중에서 사용자 질문과 **직접 관련된 내용만 추려서** 비교해줘.
# - 질문과 관련 없는 일반 설명은 제거하고, **질문 키워드와 의미상 연결된 내용**만 남겨줘.
# - 기준별로 1~2문장씩 간결하게 요약해줘.
# - 문서에 **직접 언급된 내용만** 바탕으로 답변할 것
# - 문서에 없는 정보는 “문서에 언급되지 않음”이라고 말할 것
# - 관련 내용이 없는 기준은 "관련 내용 없음"이라고 표시해줘.

# [기준별 응답들]
# """
#     for std, resp in filtered_responses.items():
#         refined_resp = extract_relevant_sentences(resp, keywords)
#         prompt += f"\n### {std} 응답:\n{refined_resp}\n"

#     messages = [
#         SystemMessage(content="너는 ESG 공시 기준 전문가야. 사용자의 질문에 따라 기준별 응답을 요약·비교해줘."),
#         HumanMessage(content=prompt)
#     ]
#     return llm(messages).content


# ✅ Streamlit 앱 실행 함수
# def run_global_chatbot():
#     st.header("ESG Chatbot")
#     query = st.text_input("질문을 입력하세요 (예: 온실가스 배출을 어떻게 공시해야 하나요?)")

#     if query:
#         with st.spinner("질문 분석 중: 관련 기준서를 식별하는 중입니다..."):
#             selected = classify_relevant_standards(query)
#             st.markdown(f"✅ 선택된 기준서: {', '.join(selected)}")

#         with st.spinner("기준별 문서 검색 및 응답 생성 중..."):
#             responses, references = generate_response_for_each_standard(query, selected)

#         with st.expander("📄 기준별 상세 응답 보기"):
#             for std, resp in responses.items():
#                 st.markdown(f"### {std} 응답")
#                 st.write(resp)

#         with st.expander("📚 인용된 문서 출처 보기"):
#             for std, refs in references.items():
#                 st.markdown(f"## {std} 문서 인용")
#                 for ref in refs:
#                     cleaned_text = clean_text(ref['content'])
#                     content_html = f"""
#                     <div style='margin-bottom:1.5em; padding:1em; border:1px solid #ccc; border-radius:10px; background-color:#f9f9f9'>
#                         <p><b>📌 소스:</b> {ref.get("source", "없음")}</p>
#                         <pre style='white-space: pre-wrap; font-size: 14px; line-height: 1.5;'>{cleaned_text}</pre>
#                     </div>
#                     """
#                     st.markdown(content_html, unsafe_allow_html=True)


def run_eng_chatbot_app():
    import os
    import re
    import streamlit as st
    from io import BytesIO
    import pythoncom
    from docx import Document
    import tempfile
    from docx2pdf import convert

    def remove_control_characters(text: str) -> str:
        return re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

    # st.set_page_config(page_title="ESG 공시 비교 챗봇", layout="wide")
    st.title("🌐G-Mate")

    user_input = st.text_input("질문을 입력하세요 (예: 온실가스 배출을 어떻게 공시해야 하나요?)")

    # if st.button("입력") and user_input:
    #     with st.spinner("질문 분석 중: 관련 기준서를 식별하는 중입니다..."):
    #         selected = classify_relevant_standards(user_input)
    #         st.markdown(f"✅ 선택된 기준서: {', '.join(selected) if selected else '없음'}")

    #     with st.spinner("기준별 문서 검색 및 응답 생성 중..."):
    #         responses, references = generate_response_for_each_standard(user_input, selected)
        
    #     st.session_state.responses = responses
    #     st.session_state.references = references
    #     st.session_state.user_input = user_input

    if st.button("입력") and user_input:
        with st.spinner("질문 분석 중: 관련 기준서를 식별하는 중입니다..."):
            # ✅ 비교 질문 여부 판단
            is_comparative = is_comparative_question_via_llm(user_input)
            st.session_state.is_comparative = is_comparative
            if is_comparative:
                st.info("📌 이 질문은 기준서 간 비교 질문입니다.")

            selected = classify_relevant_standards(user_input)
            st.markdown(f"✅ 선택된 기준서: {', '.join(selected) if selected else '없음'}")

        with st.spinner("기준별 문서 검색 및 응답 생성 중..."):
            responses, references = generate_response_for_each_standard(user_input, selected)

        st.session_state.responses = responses
        st.session_state.references = references
        st.session_state.user_input = user_input



    if "user_input" in st.session_state:
        responses = st.session_state.responses
        references = st.session_state.references
        user_input = st.session_state.user_input

        # st.header(f"📄 기준별 상세 응답 보기")
        # for std, resp in responses.items():
        #     st.markdown(f"### {std} 응답")
        #     st.write(resp)
        for std, resp in responses.items():
            for ref in references.get(std, []):
                raw_english = ref['content'].strip()

                # ✅ 영어 문단을 한국어로 번역
                translated = translate_to_korean(raw_english).strip()
                # ✅ 유사도 기반 비교
                if not is_similar_sentence(translated, resp):
                    print(f"[⚠️] {std} 응답에 인용된 문단이 포함되지 않음")


        if responses:
            st.header("📄 기준별 상세 응답 보기")
            for std, resp in responses.items():
                st.markdown(f"### {std} 응답")
                st.write(resp)
        else:
            st.header("📄 기준별 상세 응답 없음")

        with st.expander("📚 인용된 문서 출처 보기"):
            for std, refs in references.items():
                st.markdown(f"## {std} 문서 인용")
                for ref in refs:
                    cleaned_text = clean_text(ref['content'])
                    content_html = f"""
                    <div style='margin-bottom:1.5em; padding:1em; border:1px solid #ccc; border-radius:10px; background-color:#f9f9f9'>
                        <p><b>📌 소스:</b> {ref.get("source", "없음")}</p>
                        <pre style='white-space: pre-wrap; font-size: 14px; color: auto; line-height: 1.5;'>{cleaned_text}</pre>
                    </div>
                    """
                    st.markdown(content_html, unsafe_allow_html=True)

        # Word 문서 생성
        doc = Document()
        doc.add_heading('ESG 공시 비교 챗봇 결과', 0)
        doc.add_heading('질문', level=1)
        doc.add_paragraph(user_input)
        doc.add_heading('📄 기준별 응답', level=1)
        # if "summary" in responses: ##추가
        #     st.subheader("🧩 기준서 간 응답 요약")
        #     st.write(responses["summary"])

        for std, resp in responses.items():
            doc.add_heading(std, level=2)
            doc.add_paragraph(resp)
        doc.add_heading('📚 인용 문서', level=1)
        for std, refs in references.items():
            doc.add_heading(std, level=2)
            for ref in refs:
                doc.add_paragraph(f"출처: {ref['source']}")
                safe_text = remove_control_characters(clean_text(ref['content']))
                doc.add_paragraph(safe_text)

        # Word → BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        st.download_button(
            label="⬇️ Word 파일 다운로드",
            data=doc_io,
            file_name="esg_답변기록.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # PDF 다운로드
        def convert_docx_to_pdf(doc: Document) -> BytesIO:
            pythoncom.CoInitialize()
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "temp.docx")
                pdf_path = os.path.join(tmpdir, "temp.pdf")
                doc.save(docx_path)
                convert(docx_path, pdf_path)
                with open(pdf_path, "rb") as f:
                    return BytesIO(f.read())

        # ✅ PDF 다운로드 버튼 처리
        try:
            pdf_bytes = convert_docx_to_pdf(doc)
            st.download_button(
                label="⬇️ PDF 파일 다운로드",
                data=pdf_bytes,
                file_name="esg_답변기록.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.warning("⚠️ PDF 저장 중 오류가 발생했습니다. Word 파일은 정상 저장됩니다.")
            st.text(str(e))
